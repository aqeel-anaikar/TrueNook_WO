import pandas as pd
from docx import Document
from datetime import datetime
import os

# =========================
# CONFIGURATION
# =========================
INPUT_FILE = "trainers.xlsx"      # change to trainers.csv if needed
TEMPLATE_FILE = "template.docx"   # your work order template
OUTPUT_FOLDER = "generated_work_orders"

TDS_PERCENT = 0.10

# Create output folder
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


# =========================
# LOAD DATA
# =========================
def load_data(file_path):
    if file_path.endswith(".csv"):
        return pd.read_csv(file_path)
    else:
        return pd.read_excel(file_path)


# =========================
# TEXT REPLACEMENT FUNCTION
# =========================
def replace_text(doc, key, value):
    value = str(value)

    def replace_in_paragraph(para):
        """Replace placeholder in paragraph, handling split runs"""
        runs = para.runs
        if not runs:
            return
        # Join all text to find placeholder
        full_text = ''.join([r.text for r in runs])
        if key not in full_text:
            return
        # Create new text with replacement
        new_text = full_text.replace(key, value)
        # Clear all runs
        for i in range(len(runs) - 1, -1, -1):
            r = runs[i]._element
            r.getparent().remove(r)
        # Add new run with replaced text
        para.add_run(new_text)

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)


# =========================
# CALCULATION FUNCTION
# =========================
def calculate(total_days, per_day):
    # Split logic
    if total_days >= 40:
        days1 = 20
        days2 = total_days - 20
    else:
        days1 = total_days // 2
        days2 = total_days - days1

    total1 = days1 * per_day
    total2 = days2 * per_day

    tds1 = total1 * TDS_PERCENT
    tds2 = total2 * TDS_PERCENT

    final1 = total1 - tds1
    final2 = total2 - tds2

    return {
        "days1": days1,
        "days2": days2,
        "total1": total1,
        "total2": total2,
        "tds1": tds1,
        "tds2": tds2,
        "final1": final1,
        "final2": final2
    }


# =========================
# GENERATE DOCUMENT
# =========================
def generate_work_order(row):
    doc = Document(TEMPLATE_FILE)

    name = str(row['name'])
    aadhar = str(row['aadhar'])
    total_days = int(row['total_days'])
    
    # Parse and format dates (handle datetime strings like '2026-04-15 00:00:00')
    try:
        from_dt = pd.to_datetime(row['from_date'], errors='coerce')
        from_date = from_dt.strftime("%d-%b-%Y") if pd.notna(from_dt) else str(row['from_date'])
    except Exception:
        from_date = str(row['from_date'])
    
    try:
        to_dt = pd.to_datetime(row['to_date'], errors='coerce')
        to_date = to_dt.strftime("%d-%b-%Y") if pd.notna(to_dt) else str(row['to_date'])
    except Exception:
        to_date = str(row['to_date'])
    
    per_day = float(row['remuneration'])

    today = datetime.now().strftime("%d-%b-%Y")

    calc = calculate(total_days, per_day)

    # =========================
    # REPLACEMENTS
    # =========================
    replace_text(doc, "{{DATE}}", today)
    replace_text(doc, "{{NAME}}", name)
    replace_text(doc, "{{AADHAR}}", aadhar)
    replace_text(doc, "{{FROM_DATE}}", from_date)
    replace_text(doc, "{{TO_DATE}}", to_date)
    replace_text(doc, "{{PER_DAY}}", f"₹{per_day:,.0f}")

    replace_text(doc, "{{DAYS1}}", str(int(calc["days1"])))
    replace_text(doc, "{{DAYS2}}", str(int(calc["days2"])))

    replace_text(doc, "{{TOTAL1}}", f"₹{int(calc['total1']):,}")
    replace_text(doc, "{{TOTAL2}}", f"₹{int(calc['total2']):,}")

    replace_text(doc, "{{TDS1}}", f"{calc['tds1']:,.0f}")
    replace_text(doc, "{{TDS2}}", f"{calc['tds2']:,.0f}")

    replace_text(doc, "{{FINAL1}}", f"{calc['final1']:,.0f}")
    replace_text(doc, "{{FINAL2}}", f"{calc['final2']:,.0f}")

    # =========================
    # FORMAT ALL PARAGRAPHS (Font: Times New Roman, Size: 12)
    # =========================
    from docx.shared import Pt
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # Format table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    # =========================
    # SAVE FILE
    # =========================
    safe_name = name.replace(" ", "_")
    output_path = os.path.join(OUTPUT_FOLDER, f"Work_Order_{safe_name}.docx")

    doc.save(output_path)

    print(f"✅ Generated: {output_path}")


# =========================
# MAIN EXECUTION
# =========================
def main():
    print("🚀 Starting Work Order Automation...")

    df = load_data(INPUT_FILE)

    required_columns = ['name', 'aadhar', 'total_days', 'from_date', 'to_date', 'remuneration']

    for col in required_columns:
        if col not in df.columns:
            print(f"❌ Missing column: {col}")
            return

    for index, row in df.iterrows():
        try:
            generate_work_order(row)
        except Exception as e:
            print(f"❌ Error in row {index}: {e}")

    print("🎉 All Work Orders Generated Successfully!")


# =========================
# RUN
# =========================
if __name__ == "__main__":
    main()